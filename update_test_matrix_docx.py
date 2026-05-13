from docx import Document
from docx.shared import Pt

path = r"D:\University\EduLearn_Platform\Dashboard Documents\EduLearn_System_Tests_Plan.docx"
doc = Document(path)

def add_title(text, size=16):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)

def add_subtitle(text, size=13):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)

doc.add_page_break()
add_title('Ready-to-Run Test Matrix')
doc.add_paragraph('Scope: EduLearn Dashboard (Frontend + Backend) for Super Admin, School Admin, and Branch Admin.')
doc.add_paragraph('Columns: Test ID, Module, Role, Precondition, Steps, Expected Result, Priority, Type.')

headers = ["Test ID", "Module", "Role", "Precondition", "Steps", "Expected Result", "Priority", "Type"]
rows = [
    ["TC-AUTH-001", "Authentication", "All", "User account exists", "Open login page > enter valid credentials > click Login", "User is redirected to role-appropriate dashboard", "P0", "E2E"],
    ["TC-AUTH-002", "Authentication", "All", "User account exists", "Login with invalid password", "Validation error shown; no session created", "P0", "Negative"],
    ["TC-AUTH-003", "Security", "Branch Admin", "Account has temporary password", "Login with temp password", "User redirected to force-password-change page", "P0", "E2E"],
    ["TC-RBAC-001", "Authorization", "Super Admin", "Super admin logged in", "Open /super-admin", "Page loads successfully", "P0", "Backend+UI"],
    ["TC-RBAC-002", "Authorization", "School Admin", "School admin logged in", "Try open /super-admin", "Access denied (403 or redirect)", "P0", "Negative"],
    ["TC-RBAC-003", "Authorization", "Branch Admin", "Branch admin logged in", "Try open other school records by URL tampering", "Access denied; no cross-school data exposure", "P0", "Security"],
    ["TC-DASH-001", "Dashboard", "School Admin", "Has school data", "Open /dashboard", "KPI cards load with school-scoped counts", "P1", "UI+Data"],
    ["TC-DASH-002", "Dashboard", "Branch Admin", "Has branch data", "Open /dashboard", "Only branch/school scoped data displayed", "P0", "Data Isolation"],
    ["TC-DASH-003", "AI Insight", "School Admin", "AI service reachable", "Click refresh insight / wait async load", "Insight loads or graceful error message appears", "P2", "Async"],
    ["TC-STU-001", "Students CRUD", "School Admin", "Logged in", "Create new student with valid payload", "Student created and visible in list", "P0", "Backend+UI"],
    ["TC-STU-002", "Students CRUD", "School Admin", "Existing student", "Update student profile", "Changes persist and return in API/list", "P1", "Backend+UI"],
    ["TC-STU-003", "Students Import", "School Admin", "Valid CSV template", "Import CSV", "Rows imported; success/fail counts shown", "P1", "Import"],
    ["TC-STU-004", "Students Security", "School Admin", "Another school student ID known", "Call update/delete endpoint with foreign student ID", "Forbidden (403)", "P0", "Security"],
    ["TC-TCH-001", "Teachers CRUD", "School Admin", "Logged in", "Create teacher with assignments", "Teacher created with assignment records", "P0", "Backend+UI"],
    ["TC-TCH-002", "Teachers Import", "School Admin", "Valid XLSX template", "Import XLSX", "Teacher records imported; codes generated", "P1", "Import"],
    ["TC-CLS-001", "Classes", "School Admin", "Logged in", "Create class/section", "Class appears in classes list", "P1", "Backend+UI"],
    ["TC-SUB-001", "Subjects", "School Admin", "Logged in", "Create subject", "Subject appears and can be linked to classes", "P1", "Backend+UI"],
    ["TC-ASN-001", "Assignments", "School Admin", "Classes+teachers+subjects exist", "Create assignment", "Assignment saved and listed", "P1", "Backend+UI"],
    ["TC-RPT-001", "Reports List", "School Admin", "Student data exists", "Open /reports and load list", "Class/Student report data loads correctly", "P1", "Data"],
    ["TC-RPT-002", "At-Risk", "School Admin", "Mixed performance data", "Open at-risk report", "Only students matching risk criteria are returned", "P1", "Logic"],
    ["TC-NOTIF-001", "Notifications", "Super Admin", "Logged in", "Create global notification", "Notification persisted and broadcasted", "P0", "Realtime"],
    ["TC-NOTIF-002", "Notifications", "School Admin", "Notification exists", "Open dropdown and mark as read", "Read status updated in UI and backend", "P1", "UI+API"],
    ["TC-NOTIF-003", "Notifications", "All", "Reverb running", "Trigger broadcast event", "Role/school channel receives correct message", "P0", "Realtime"],
    ["TC-SUP-001", "Support", "School Admin", "Logged in", "Submit support ticket", "Ticket created with open status", "P1", "Backend+UI"],
    ["TC-SUP-002", "Support", "Super Admin", "Ticket exists", "Reply and close ticket", "Ticket status and messages updated", "P1", "Workflow"],
    ["TC-SET-001", "Preferences", "School Admin", "Logged in", "Toggle dark/light mode", "Theme persists after refresh/login", "P1", "UI+Session"],
    ["TC-SET-002", "Preferences", "School Admin", "Logged in", "Switch language ar/en", "Locale changes and persists", "P1", "UI+Session"],
    ["TC-BR-001", "Branches", "School Admin", "Logged in", "Create branch request", "Branch saved with pending status", "P1", "Workflow"],
    ["TC-BR-002", "Branches", "Super Admin", "Pending branch exists", "Approve branch school", "Status changes to active and branch admin can proceed", "P1", "Workflow"],
    ["TC-BR-003", "Branch Permissions", "School Admin", "Branch admin exists", "Grant/revoke permissions", "Permissions saved and reflected", "P1", "RBAC"],
    ["TC-SEC-001", "CSRF", "All", "Logged in", "Submit protected form without token", "Request rejected", "P0", "Security"],
    ["TC-SEC-002", "Validation", "All", "Logged in", "Submit invalid payloads on key forms", "Server-side validation blocks invalid data", "P0", "Security"],
    ["TC-EXP-001", "Export", "School Admin", "Students/teachers data exists", "Export CSV/PDF", "Files generated with correct scoped data", "P1", "Export"],
    ["TC-PERF-001", "Performance", "All", "Seeded realistic data", "Load dashboard and reports", "Response times acceptable and no server errors", "P2", "Performance"],
]

table = doc.add_table(rows=1, cols=len(headers))
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r in rows:
    cells = table.add_row().cells
    for i, val in enumerate(r):
        cells[i].text = val

doc.add_paragraph('')
add_subtitle('Execution Checklist')
for item in [
    'Prepare dedicated test accounts for each role.',
    'Reset database or use a known seeded snapshot before each full run.',
    'Run smoke tests first (AUTH, RBAC, Dashboard, Notifications).',
    'Run P0 tests before every deployment.',
    'Capture evidence (screenshot + API response) for failed cases.',
    'Log defects with: Test ID, steps, expected, actual, environment, severity.'
]:
    doc.add_paragraph('- ' + item)

doc.save(path)
print('UPDATED:', path)
