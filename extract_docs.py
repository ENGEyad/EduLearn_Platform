from docx import Document
import os

base = r'd:\University\EduLearn_Platform\تحليل عمليات تسجيل الدخول وتهيئة المدرسة وعمليات السوبر ادمن'
output_dir = r'd:\University\EduLearn_Platform'

files = [
    ('AI_Agent_Prompt_EduLearn_School_Management_System_Development_(Laravel).docx', 'laravel_ai_agent_prompt.md'),
    ('Reconciled_Requirements_EduLearn_School_Management_System.docx', 'reconciled_requirements.md'),
    ('Updated_Design_Document_EduLearn_School_Management_System.docx', 'updated_design_document.md'),
]

for docx_file, output_name in files:
    path = os.path.join(base, docx_file)
    doc = Document(path)
    
    with open(os.path.join(output_dir, output_name), 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            f.write(para.text + '\n')
        
        # Also extract tables
        for i, table in enumerate(doc.tables):
            f.write(f'\n--- Table {i+1} ---\n')
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                f.write(' | '.join(cells) + '\n')
    
    print(f'Extracted: {output_name}')

print('Done!')
